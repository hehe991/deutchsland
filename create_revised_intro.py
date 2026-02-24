# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('INTRODUCTION (Background)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Paragraph 1: Introduction (Broad to Narrow)
para1 = '''Artificial intelligence has emerged as an important development in healthcare, with applications ranging from diagnostic imaging and clinical decision support to personalized medicine and patient management (1-3). The introduction of generative AI, particularly large language models such as ChatGPT, has accelerated this transformation, offering significant opportunities for medical education, clinical training, and healthcare delivery (4, 5). While these technologies present potential benefits for enhancing learning experiences and improving healthcare outcomes, they also raise concerns regarding accuracy, bias, and reduced critical thinking skills among trainees (6, 7). Medical education institutions worldwide are increasingly recognizing the need to prepare future physicians for an AI-driven healthcare environment, leading to calls for curriculum reform and integration of AI literacy into undergraduate medical training (2, 8). Pakistan's National Health Vision 2016-2025 highlights the importance of digital health technologies and emphasizes the need for a workforce capable of using technological advancements to improve healthcare outcomes (9). However, successful integration of AI into medical practice depends not only on technological infrastructure but also on the readiness, ethical awareness, and willingness of future healthcare professionals to adopt these innovations. Despite growing global interest, there remains limited research examining medical students' preparedness for AI adoption in Pakistan, particularly within the context of local medical education systems.'''

p1 = doc.add_paragraph(para1)
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Paragraph 2: Literature Review (Current Knowledge)
para2 = '''Readiness for change is a concept that includes cognitive, affective, and behavioral components influencing individuals' capacity to adopt new technologies (10). In the context of AI in medical education, readiness refers to students' knowledge, skills, attitudes, and ethical preparedness to engage with AI technologies in their future practice. The Medical Artificial Intelligence Readiness Scale for Medical Students (MAIRS-MS) was developed as a validated tool to assess readiness across four domains: cognition, ability, vision, and ethics (11). Recent studies using this scale in Saudi Arabia reported moderate levels of AI readiness among medical students, with mean scores of 62 out of 110, indicating basic understanding but substantial room for improvement (12). Behavioral intention, an important predictor of actual technology adoption, has been studied through theoretical frameworks such as the Theory of Planned Behavior and the Technology Acceptance Model (13, 14). Research from China showed that medical students' attitudes significantly mediate the relationship between perceived usefulness, perceived ease of use, and behavioral intentions toward AI tool adoption (15). Ethical considerations surrounding AI in medical education have received increasing attention, with researchers highlighting concerns about data privacy, algorithmic bias, informed consent, and transparency of AI decision-making processes (16, 17). A comprehensive scoping review found that bias and privacy were the most frequently mentioned ethical concerns, appearing in 74% and 65% of reviewed articles respectively (18).'''

p2 = doc.add_paragraph(para2)
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Paragraph 3: Gaps in Literature
para3 = '''While studies from developed nations and Gulf countries have explored readiness, ethical concerns, and behavioral intentions, similar research in South Asian contexts, particularly Pakistan, remains limited (19, 20). Recent research from Pakistan found that only 21.3% of healthcare professionals report good familiarity with AI concepts, despite 64.1% demanding AI integration into healthcare systems (19, 20). Furthermore, no study has comprehensively examined the relationships between readiness, ethical concerns, and behavioral intentions among Pakistani medical students using validated instruments. Most existing research from Pakistan focuses on general awareness among healthcare professionals rather than systematic assessment of medical students' readiness, ethical perceptions, and behavioral intentions. Additionally, validated instruments such as MAIRS-MS have not been adapted or used in the Pakistani context to measure AI readiness systematically.'''

p3 = doc.add_paragraph(para3)
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Rationale Section (Separate)
doc.add_paragraph()
rationale_title = doc.add_heading('Rationale:', level=2)

rationale_text = '''This study addresses these identified gaps by providing a comprehensive assessment of AI readiness, ethical concerns, and behavioral intentions among medical students at University College of Medicine and Dentistry, Pakistan. The findings will provide baseline data for curriculum planners to design evidence-based AI education modules, support policy implementation of Pakistan's digital health agenda, and contribute perspectives from a developing nation to the existing literature on AI in medical education. Understanding these constructs will help identify specific areas requiring intervention and prepare medical graduates for an increasingly AI-integrated healthcare environment.'''

r_para = doc.add_paragraph(rationale_text)
r_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# References Section
doc.add_page_break()
ref_title = doc.add_heading('REFERENCES', level=1)
ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

references = [
    "Jiang F, Jiang Y, Zhi H, et al. Artificial intelligence in healthcare: past, present and future. Stroke Vasc Neurol. 2017;2:230-243.",
    "Zhang P, Kamel Boulos MN. Generative AI in Medicine and Healthcare: Promises, Opportunities and Challenges. Future Internet. 2023;15(9):286.",
    "Davenport T, Kalakota R. The potential for artificial intelligence in healthcare. Future Healthc J. 2019;6(2):94-98.",
    "Li Y, Li J. Generative artificial intelligence in medical education: way to solve the problems. Postgrad Med J. 2024;100(1181):203-204.",
    "Sallam M. ChatGPT Utility in Healthcare Education, Research, and Practice: Systematic Review. Healthcare. 2023;11(6):887.",
    "Floridi L, Chiriatti M. GPT-3: Its nature, scope, limits, and consequences. Minds Mach. 2020;30(4):681-694.",
    "Bender EM, Gebru T, McMillan-Major A, Shmitchell S. On the dangers of stochastic parrots: Can language models be too big? Proc 2021 ACM FaccT Conf. 2021:610-623.",
    "Ngiam KY, Khor IW. Big data and machine learning algorithms for health-care delivery. Lancet Oncol. 2019;20(5):e262-e273.",
    "Ministry of National Health Services, Regulations and Coordination. National Health Vision Pakistan 2016-2025. Islamabad: Government of Pakistan; 2016.",
    "Yun J, Kim DJ. Understanding readiness for change: A literature review and a new conceptual framework. J Organ Change Manag. 2021;34(1):1-22.",
    "Karaca O, Caliskan SA, Demir K. Medical artificial intelligence readiness scale for medical students (MAIRS-MS) - development, validity and reliability study. BMC Med Educ. 2021;21:112.",
    "Almalki M, Alkhamis MA, Khairallah FM, Choukou MA. Perceived artificial intelligence readiness in medical and health sciences education: a survey study of students in Saudi Arabia. BMC Med Educ. 2025;25:439.",
    "Ajzen I. The theory of planned behavior. Organ Behav Hum Decis Process. 1991;50(2):179-211.",
    "Fishbein M, Ajzen I. Predicting and changing behavior: The reasoned action approach. New York: Psychology Press; 2010.",
    "Liu F, Chang X, Zhu Q, Huang Y, Li Y, Wang H. Assessing clinical medicine students' acceptance of large language model: based on technology acceptance model. BMC Med Educ. 2024;24:1251.",
    "Luxton DD. Ethical implications of artificial intelligence in education and training. AI Ethics. 2022;2:181-189.",
    "Stahl BC, Wright D, Friedewald M. The ethics of AI in medical contexts. AI Soc. 2021;36(1):1-10.",
    "Itani A, Gronseth SL, Musaad S, Nguyen T, Mirabile Y, Beech BM. Ethical considerations for teaching with artificial intelligence: a scoping review in medical education settings. Int J Educ Technol High Educ. 2025;22(1):68.",
    "Naseer MA, Saeed S, Afzal A, Ali S, Malik MGR. Navigating the integration of artificial intelligence in the medical education curriculum: a mixed-methods study exploring the perspectives of medical students and faculty in Pakistan. BMC Med Educ. 2025;25:273.",
    "Umer M, Naveed A, Maryam Q, Malik AR, Bashir N, Kandel K. Investigating awareness of artificial intelligence in healthcare among medical students and professionals in Pakistan: a cross-sectional study. Ann Med Surg. 2024;86(5):2606-2611.",
    "Chen J, Wang Y, Wang X. Factors affecting the behavioral intention to use AI-powered learning tools in medical education. BMC Med Educ. 2024;24:112."
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(ref)

doc.save(r'C:\Users\haseeb\Downloads\Introduction_AI_Adoption_UCMD.docx')
print('Document updated successfully!')
print('\nChanges made:')
print('1. Simplified intro language (more natural for undergraduate student)')
print('2. Fixed "Al" -> "AI" typos')
print('3. Rationale kept as separate section')
print('4. Removed repetition between paragraph 3 and rationale')
print('5. Paragraph 3 now focuses on gaps only')
print('6. Rationale focuses on study significance and contributions')
